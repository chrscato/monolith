# billing/logic/postprocess/utils/eobr_generator.py

import logging
import os
import tempfile
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime, date
import re

try:
    from docx import Document
except ImportError:
    raise ImportError("python-docx is required. Install with: pip install python-docx")

logger = logging.getLogger(__name__)

class EOBRGenerator:
    """
    Generator for EOBR (Explanation of Bill Review) documents.
    Creates DOCX files from templates with bill data.
    """
    
    def __init__(self, template_path: Path = None):
        """
        Initialize the EOBR generator.
        
        Args:
            template_path: Path to the EOBR template file
        """
        self.template_path = self._resolve_template_path(template_path)
        self._validate_template()
        
        logger.info(f"EOBR Generator initialized with template: {self.template_path}")
    
    def _resolve_template_path(self, template_path: Path = None) -> Path:
        """
        Resolve the template path with multiple fallback options.
        
        Args:
            template_path: Optional explicit template path
            
        Returns:
            Resolved template path
        """
        if template_path and template_path.exists():
            return template_path
        
        # Get project root (4 levels up from this file)
        project_root = Path(__file__).resolve().parents[4]
        
        # Try multiple possible locations
        possible_paths = [
            # Explicit path if provided
            template_path,
            # Billing templates directory
            project_root / "billing" / "templates" / "EOBR Template.docx",
            # Project root
            project_root / "EOBR Template.docx",
            # Current directory
            Path("EOBR Template.docx"),
            # Same directory as this script
            Path(__file__).parent / "EOBR Template.docx",
        ]
        
        for path in possible_paths:
            if path and path.exists():
                logger.info(f"Found template at: {path}")
                return path
        
        # If no template found, use the preferred location
        preferred_path = project_root / "billing" / "templates" / "EOBR Template.docx"
        logger.warning(f"Template not found. Expected location: {preferred_path}")
        return preferred_path
    
    def _validate_template(self):
        """Validate that the template file exists and is readable."""
        if not self.template_path.exists():
            raise FileNotFoundError(
                f"EOBR template not found at {self.template_path}. "
                f"Please create the template file with required placeholders."
            )
        
        try:
            # Try to open the template to validate it's a valid DOCX
            doc = Document(str(self.template_path))
            logger.debug(f"Template validated successfully: {len(doc.paragraphs)} paragraphs")
        except Exception as e:
            raise ValueError(f"Invalid DOCX template: {str(e)}")
    
    def format_currency(self, amount: Any) -> str:
        """
        Format currency amount.
        
        Args:
            amount: Amount to format (float, int, string, or None)
            
        Returns:
            Formatted currency string ($X.XX)
        """
        if amount is None:
            return "$0.00"
        
        try:
            # Handle string amounts
            if isinstance(amount, str):
                # Remove currency symbols and commas
                clean_amount = re.sub(r'[$,]', '', amount.strip())
                amount = float(clean_amount)
            
            # Convert to float and format
            amount = float(amount)
            return f"${amount:.2f}"
            
        except (ValueError, TypeError):
            logger.warning(f"Could not format currency amount: {amount}")
            return "$0.00"
    
    def format_date(self, date_value: Any) -> str:
        """
        Format date value to MM/DD/YYYY.
        
        Args:
            date_value: Date to format (string, date, datetime, or None)
            
        Returns:
            Formatted date string (MM/DD/YYYY)
        """
        if not date_value:
            return ""
        
        try:
            if isinstance(date_value, str):
                date_str = date_value.strip()
                
                # Handle date ranges (take first date)
                if ' - ' in date_str:
                    date_str = date_str.split(' - ')[0].strip()
                
                # Try different date formats
                date_formats = [
                    '%Y-%m-%d',      # 2024-01-15
                    '%m/%d/%Y',      # 01/15/2024
                    '%m/%d/%y',      # 01/15/24
                    '%m-%d-%Y',      # 01-15-2024
                    '%m-%d-%y',      # 01-15-24
                ]
                
                for fmt in date_formats:
                    try:
                        parsed_date = datetime.strptime(date_str, fmt).date()
                        return parsed_date.strftime('%m/%d/%Y')
                    except ValueError:
                        continue
                
                # If no format matches, return as-is
                logger.warning(f"Could not parse date format: {date_value}")
                return str(date_value)
            
            elif isinstance(date_value, (date, datetime)):
                return date_value.strftime('%m/%d/%Y')
            
            else:
                return str(date_value)
                
        except Exception as e:
            logger.warning(f"Error formatting date {date_value}: {str(e)}")
            return str(date_value) if date_value else ""
    
    def prepare_bill_data(self, bill: Dict[str, Any]) -> Dict[str, str]:
        """
        Prepare bill data for EOBR template replacement.
        
        Args:
            bill: Bill dictionary with all required data
            
        Returns:
            Dictionary of placeholder mappings
        """
        try:
            line_items = bill.get('line_items', [])
            
            # Calculate total paid (sum of allowed amounts)
            total_paid = 0.0
            for item in line_items:
                allowed_amount = item.get('allowed_amount', 0)
                try:
                    total_paid += float(allowed_amount) if allowed_amount is not None else 0.0
                except (ValueError, TypeError):
                    logger.warning(f"Invalid allowed_amount: {allowed_amount}")
                    continue
            
            # Prepare header data
            data = {
                # Header section
                'PatientName': str(bill.get('PatientName', '')),
                'dob': self.format_date(bill.get('Patient_DOB')),
                'process_date': datetime.now().strftime('%m/%d/%Y'),
                'order_no': str(bill.get('Order_ID', '')),
                'provider_ref': str(bill.get('FileMaker_Record_Number', '')),
                'doi': self.format_date(bill.get('Patient_Injury_Date')),
                
                # Provider section
                'TIN': str(bill.get('provider_tin', '')),
                'NPI': str(bill.get('provider_npi', '')),
                'billing_name': str(bill.get('provider_billing_name', '')),
                'billing_address1': str(bill.get('provider_billing_address1', '')),
                'billing_address2': str(bill.get('provider_billing_address2', '')),
                'billing_city': str(bill.get('provider_billing_city', '')),
                'billing_state': str(bill.get('provider_billing_state', '')),
                'billing_zip': str(bill.get('provider_billing_postal_code', '')),
                
                # Footer
                'total_paid': self.format_currency(total_paid)
            }
            
            # Prepare line items (up to 6 slots)
            for i in range(1, 7):  # dos1 through dos6
                if i <= len(line_items):
                    item = line_items[i-1]
                    
                    # Get amounts
                    allowed_amount = 0.0
                    charge_amount = 0.0
                    
                    try:
                        allowed_amount = float(item.get('allowed_amount', 0)) if item.get('allowed_amount') is not None else 0.0
                        charge_amount = float(item.get('charge_amount', 0)) if item.get('charge_amount') is not None else 0.0
                    except (ValueError, TypeError):
                        logger.warning(f"Invalid amounts in line item {i}")
                    
                    # Determine paid amount and reason code
                    if allowed_amount == 0:
                        reason_code = "125"  # Denied
                        paid_amount = 0.0
                    elif allowed_amount < charge_amount:
                        reason_code = "85"   # Reduced payment
                        paid_amount = allowed_amount
                    else:
                        reason_code = "85"   # Full payment
                        paid_amount = allowed_amount
                    
                    # Format modifier
                    modifier = item.get('modifier', '')
                    if modifier and modifier.strip():
                        modifier = modifier.strip()
                    else:
                        modifier = ''
                    
                    data.update({
                        f'dos{i}': self.format_date(item.get('date_of_service')),
                        f'pos{i}': str(item.get('place_of_service', '11')),
                        f'cpt{i}': str(item.get('cpt_code', '')),
                        f'modifier{i}': modifier,
                        f'units{i}': str(item.get('units', 1)),
                        f'charge{i}': self.format_currency(charge_amount),
                        f'alwd{i}': self.format_currency(allowed_amount),
                        f'paid{i}': self.format_currency(paid_amount),
                        f'code{i}': reason_code
                    })
                else:
                    # Empty slots
                    data.update({
                        f'dos{i}': '',
                        f'pos{i}': '',
                        f'cpt{i}': '',
                        f'modifier{i}': '',
                        f'units{i}': '',
                        f'charge{i}': '',
                        f'alwd{i}': '',
                        f'paid{i}': '',
                        f'code{i}': ''
                    })
            
            logger.debug(f"Prepared EOBR data for bill {bill.get('bill_id', 'unknown')}")
            return data
            
        except Exception as e:
            logger.error(f"Error preparing bill data: {str(e)}")
            raise
    
    def replace_placeholders_in_text(self, text: str, data: Dict[str, str]) -> str:
        """
        Replace placeholders in text with actual data.
        
        Args:
            text: Text containing placeholders like <PatientName>
            data: Dictionary of placeholder mappings
            
        Returns:
            Text with placeholders replaced
        """
        if not text:
            return text
        
        def replace_placeholder(match):
            placeholder = match.group(1)
            replacement = data.get(placeholder, f"<{placeholder}>")  # Keep original if not found
            return replacement
        
        # Find all placeholders in format <placeholder>
        result = re.sub(r'<([^>]+)>', replace_placeholder, text)
        return result
    
    def replace_text_in_run(self, run, old_text: str, new_text: str):
        """
        Replace text in a run while preserving formatting.
        
        Args:
            run: Document run object
            old_text: Text to replace
            new_text: Replacement text
        """
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
    
    def replace_placeholders_in_paragraph(self, paragraph, data: Dict[str, str]):
        """
        Replace placeholders in a paragraph while preserving formatting.
        
        Args:
            paragraph: Document paragraph object
            data: Dictionary of placeholder mappings
        """
        # Check if paragraph contains placeholders
        if not paragraph.text or '<' not in paragraph.text:
            return
        
        # Find all placeholders in the paragraph
        placeholders = re.findall(r'<([^>]+)>', paragraph.text)
        
        if not placeholders:
            return
        
        # Replace each placeholder in all runs
        for placeholder in placeholders:
            placeholder_text = f'<{placeholder}>'
            replacement_text = data.get(placeholder, placeholder_text)
            
            # Go through each run and replace the placeholder
            for run in paragraph.runs:
                if placeholder_text in run.text:
                    self.replace_text_in_run(run, placeholder_text, replacement_text)
    
    def replace_placeholders_in_document(self, doc: Document, data: Dict[str, str]):
        """
        Replace placeholders throughout the entire document while preserving formatting.
        
        Args:
            doc: python-docx Document object
            data: Dictionary of placeholder mappings
        """
        try:
            # Replace in paragraphs
            for paragraph in doc.paragraphs:
                self.replace_placeholders_in_paragraph(paragraph, data)
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self.replace_placeholders_in_paragraph(paragraph, data)
            
            # Replace in headers and footers
            for section in doc.sections:
                # Header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        self.replace_placeholders_in_paragraph(paragraph, data)
                
                # Footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        self.replace_placeholders_in_paragraph(paragraph, data)
            
            logger.debug("Successfully replaced placeholders in document")
            
        except Exception as e:
            logger.error(f"Error replacing placeholders in document: {str(e)}")
            raise
    
    def generate_eobr(self, bill: Dict[str, Any], output_path: Path) -> bool:
        """
        Generate an EOBR document for a single bill.
        
        Args:
            bill: Bill dictionary with all required data including line_items
            output_path: Path where the generated EOBR should be saved
            
        Returns:
            True if successful, False otherwise
        """
        try:
            bill_id = bill.get('bill_id', bill.get('id', 'unknown'))
            logger.info(f"Generating EOBR for bill {bill_id}")
            
            # Prepare data for replacement
            data = self.prepare_bill_data(bill)
            
            # Load template
            doc = Document(str(self.template_path))
            
            # Replace placeholders
            self.replace_placeholders_in_document(doc, data)
            
            # Ensure output directory exists
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Save document
            doc.save(str(output_path))
            
            # Verify file was created
            if output_path.exists() and output_path.stat().st_size > 0:
                logger.info(f"EOBR generated successfully: {output_path}")
                return True
            else:
                logger.error(f"EOBR file not created or is empty: {output_path}")
                return False
            
        except Exception as e:
            logger.error(f"Error generating EOBR for bill {bill.get('bill_id', 'unknown')}: {str(e)}")
            return False
    
    def generate_batch_eobrs(self, 
                           bills: List[Dict[str, Any]], 
                           output_dir: Path,
                           filename_pattern: str = "EOBR_{bill_id}_{patient_name}.docx") -> List[Path]:
        """
        Generate EOBR documents for multiple bills.
        
        Args:
            bills: List of bill dictionaries
            output_dir: Directory where EOBRs should be saved
            filename_pattern: Pattern for output filenames (can use {bill_id}, {patient_name}, etc.)
            
        Returns:
            List of paths to successfully generated EOBRs
        """
        generated_files = []
        
        logger.info(f"Generating EOBRs for {len(bills)} bills")
        
        for bill in bills:
            try:
                # Prepare filename
                bill_id = bill.get('bill_id', bill.get('id', 'unknown'))
                patient_name = bill.get('PatientName', 'unknown_patient')
                
                # Clean patient name for filename (remove invalid characters)
                safe_patient_name = re.sub(r'[<>:"/\\|?*]', '_', patient_name)
                safe_patient_name = safe_patient_name.replace(' ', '_')
                
                filename = filename_pattern.format(
                    bill_id=bill_id,
                    patient_name=safe_patient_name,
                    order_id=bill.get('Order_ID', ''),
                    fm_record=bill.get('FileMaker_Record_Number', ''),
                    date=datetime.now().strftime('%Y%m%d')
                )
                
                output_path = output_dir / filename
                
                # Generate EOBR
                if self.generate_eobr(bill, output_path):
                    generated_files.append(output_path)
                    
            except Exception as e:
                logger.error(f"Error processing bill {bill.get('bill_id', 'unknown')}: {str(e)}")
                continue
        
        logger.info(f"Successfully generated {len(generated_files)} EOBRs out of {len(bills)} bills")
        return generated_files
    
    def get_template_placeholders(self) -> List[str]:
        """
        Extract all placeholders from the template document.
        
        Returns:
            List of placeholder names found in the template
        """
        try:
            doc = Document(str(self.template_path))
            placeholders = set()
            
            # Search in paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    found = re.findall(r'<([^>]+)>', paragraph.text)
                    placeholders.update(found)
            
            # Search in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text:
                                found = re.findall(r'<([^>]+)>', paragraph.text)
                                placeholders.update(found)
            
            return sorted(list(placeholders))
            
        except Exception as e:
            logger.error(f"Error extracting placeholders: {str(e)}")
            return []

def generate_eobr_documents(bills: List[Dict[str, Any]], 
                          output_dir: Path = None,
                          template_path: Path = None) -> List[Path]:
    """
    Convenience function to generate EOBR documents.
    
    Args:
        bills: List of prepared bill dictionaries
        output_dir: Output directory (defaults to temp directory)
        template_path: Path to EOBR template
        
    Returns:
        List of paths to generated EOBR files
    """
    if output_dir is None:
        output_dir = Path(tempfile.mkdtemp()) / "eobrs"
    
    generator = EOBRGenerator(template_path)
    return generator.generate_batch_eobrs(bills, output_dir)

def print_template_requirements():
    """Print the required template placeholders."""
    placeholders = {
        'Header': ['PatientName', 'dob', 'process_date', 'order_no', 'provider_ref', 'doi'],
        'Provider': ['TIN', 'NPI', 'billing_name', 'billing_address1', 'billing_address2', 
                    'billing_city', 'billing_state', 'billing_zip'],
        'Service Lines (1-6)': ['dos1-dos6', 'pos1-pos6', 'cpt1-cpt6', 'modifier1-modifier6', 
                               'units1-units6', 'charge1-charge6', 'alwd1-alwd6', 'paid1-paid6', 'code1-code6'],
        'Footer': ['total_paid']
    }
    
    print("\n" + "="*60)
    print("EOBR TEMPLATE PLACEHOLDER REQUIREMENTS")
    print("="*60)
    
    for section, items in placeholders.items():
        print(f"\n{section.upper()}:")
        for item in items:
            if '-' in item:  # Range notation
                base, range_part = item.split('-')
                start_num = base[-1]
                end_num = range_part[-1]
                print(f"  <{base[:-1]}1> through <{base[:-1]}6>")
            else:
                print(f"  <{item}>")
    
    print("\n" + "="*60)

if __name__ == "__main__":
    # Test the EOBR generator
    import logging
    
    logging.basicConfig(level=logging.INFO)
    
    # Print template requirements
    print_template_requirements()
    
    # Sample bill data for testing
    test_bill = {
        'bill_id': 'TEST_001',
        'PatientName': 'John Doe',
        'Patient_DOB': '1980-01-15',
        'Order_ID': 'ORD_12345',
        'FileMaker_Record_Number': 'FM_001',
        'Patient_Injury_Date': '2024-01-01',
        'provider_tin': '12-3456789',
        'provider_npi': '1234567890',
        'provider_billing_name': 'Test Medical Center',
        'provider_billing_address1': '123 Medical Drive',
        'provider_billing_address2': 'Suite 100',
        'provider_billing_city': 'Orlando',
        'provider_billing_state': 'FL',
        'provider_billing_postal_code': '32801',
        'line_items': [
            {
                'date_of_service': '2024-01-15',
                'place_of_service': '11',
                'cpt_code': '99213',
                'modifier': 'LT',
                'units': 1,
                'charge_amount': 150.00,
                'allowed_amount': 120.00
            },
            {
                'date_of_service': '2024-01-15',
                'place_of_service': '11',
                'cpt_code': '73610',
                'modifier': '',
                'units': 1,
                'charge_amount': 200.00,
                'allowed_amount': 180.00
            }
        ]
    }
    
    # Test generation
    try:
        generator = EOBRGenerator()
        output_path = Path("test_eobr.docx")
        success = generator.generate_eobr(test_bill, output_path)
        if success:
            print(f"✅ Test EOBR generated: {output_path}")
            
            # Show template placeholders found
            placeholders = generator.get_template_placeholders()
            if placeholders:
                print(f"📋 Template placeholders found: {placeholders}")
        else:
            print("❌ Failed to generate test EOBR")
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        print("\nMake sure you have:")
        print("1. Created 'EOBR Template.docx' with required placeholders")
        print("2. Installed python-docx: pip install python-docx")